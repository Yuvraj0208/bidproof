"""Proposal export + the export blocker (US-10, SPEC §5.7, §14).

Export REFUSES while any mandatory clause is unaddressed or any proposal
claim is contradicted. An override needs a name and a written reason and is
written to the append-only audit log. The moment the system refuses to export
an unproven proposal is the moment it stops being a demo and becomes a
product.
"""

import io
import uuid

from sqlalchemy import select

from app.core.db import org_scoped_session
from app.models import (
    AuditLog,
    Decision,
    Proposal,
    ProposalSection,
    Rule,
    Tender,
    VerdictRow,
)

# A mandatory clause is an eligibility requirement; it is "unaddressed" if we
# do not satisfy it (a gap) or cannot yet decide it (needs a human).
MANDATORY_FAMILY = "eligibility"
UNADDRESSED_VERDICTS = {"gap", "needs_human"}


async def export_blockers(org_id: uuid.UUID, tender_id: uuid.UUID) -> list[dict]:
    """Every reason the proposal may not be exported. Empty means clear."""
    blockers: list[dict] = []
    async with org_scoped_session(org_id) as session:
        mandatory = (
            await session.execute(
                select(VerdictRow, Rule)
                .join(Rule, VerdictRow.rule_id == Rule.rule_id)
                .where(VerdictRow.tender_id == tender_id)
                .where(Rule.family == MANDATORY_FAMILY)
                .where(VerdictRow.verdict.in_(UNADDRESSED_VERDICTS))
            )
        ).all()
        for verdict, rule in mandatory:
            blockers.append({
                "type": "unaddressed_mandatory_clause",
                "rule_key": rule.key,
                "verdict": verdict.verdict,
                "message": f"mandatory clause '{rule.key}' is {verdict.verdict} — "
                           "it must be addressed before export",
            })

        proposal = (
            await session.execute(
                select(Proposal).where(Proposal.tender_id == tender_id)
            )
        ).scalar_one_or_none()
        if proposal is not None:
            sections = (
                await session.execute(
                    select(ProposalSection).where(
                        ProposalSection.proposal_id == proposal.id
                    )
                )
            ).scalars().all()
            for section in sections:
                contradicted = [
                    c for c in (section.claims or []) if c["status"] == "contradicted"
                ]
                for claim in contradicted:
                    blockers.append({
                        "type": "contradicted_claim",
                        "section": section.section_tag,
                        "message": f"section '{section.section_tag}' contains a "
                                   "contradicted claim — it must be corrected "
                                   "before export",
                        "claim": claim["text"][:200],
                    })
    return blockers


async def _build_docx(org_id: uuid.UUID, tender_id: uuid.UUID) -> bytes:
    from docx import Document as Docx

    async with org_scoped_session(org_id) as session:
        tender = await session.get(Tender, tender_id)
        decision = (
            await session.execute(
                select(Decision).where(Decision.tender_id == tender_id)
            )
        ).scalar_one_or_none()
        proposal = (
            await session.execute(
                select(Proposal).where(Proposal.tender_id == tender_id)
            )
        ).scalar_one_or_none()
        sections = []
        if proposal is not None:
            sections = (
                await session.execute(
                    select(ProposalSection)
                    .where(ProposalSection.proposal_id == proposal.id)
                    .order_by(ProposalSection.position)
                )
            ).scalars().all()
        verdicts = (
            await session.execute(
                select(VerdictRow, Rule)
                .join(Rule, VerdictRow.rule_id == Rule.rule_id)
                .where(VerdictRow.tender_id == tender_id)
                .order_by(Rule.family, Rule.key)
            )
        ).all()

    doc = Docx()
    doc.add_heading(f"Proposal — {tender.title if tender else tender_id}", level=0)
    if decision is not None and decision.ev_inr is not None:
        doc.add_paragraph(
            f"Bid decision: {decision.recommendation.upper()} · "
            f"expected value ₹{float(decision.ev_inr) / 1e5:.2f} lakh"
        )

    for section in sections:
        doc.add_heading(section.section_tag.replace("_", " ").title(), level=1)
        doc.add_paragraph(section.content)

    # The compliance matrix, attached.
    doc.add_heading("Compliance Matrix", level=1)
    table = doc.add_table(rows=1, cols=4)
    header = table.rows[0].cells
    header[0].text, header[1].text = "Family", "Requirement"
    header[2].text, header[3].text = "Verdict", "Our position"
    for verdict, rule in verdicts:
        row = table.add_row().cells
        row[0].text = rule.family
        row[1].text = rule.key
        row[2].text = verdict.verdict
        row[3].text = (verdict.reason or "")[:200]

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


async def export_proposal(
    org_id: uuid.UUID,
    tender_id: uuid.UUID,
    override_name: str | None = None,
    override_reason: str | None = None,
) -> tuple[bytes | None, list[dict]]:
    """Returns (document_bytes, blockers). If there are blockers and no valid
    override, document_bytes is None. A valid override unblocks and is logged."""
    blockers = await export_blockers(org_id, tender_id)

    if blockers:
        if not (override_name and override_reason):
            return None, blockers
        async with org_scoped_session(org_id) as session:
            session.add(AuditLog(
                org_id=org_id, actor=override_name, action="export_override",
                tender_id=tender_id,
                details={"reason": override_reason,
                         "blockers": [b["type"] for b in blockers],
                         "blocker_detail": blockers},
            ))

    document = await _build_docx(org_id, tender_id)
    return document, blockers
